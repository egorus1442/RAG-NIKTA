import os
import fitz  # PyMuPDF
from docx import Document
from typing import List, Dict, Any

class TextExtractor:
    @staticmethod
    def extract_text(file_path: str) -> str:
        """Извлекает текст из файла в зависимости от его расширения"""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return TextExtractor._extract_from_pdf(file_path)
        elif file_extension == '.docx':
            return TextExtractor._extract_from_docx(file_path)
        elif file_extension == '.txt':
            return TextExtractor._extract_from_txt(file_path)
        else:
            raise ValueError(f"Неподдерживаемый формат файла: {file_extension}")
    
    @staticmethod
    def extract_text_with_metadata(file_path: str) -> Dict[str, Any]:
        """Извлекает текст с дополнительными метаданными"""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return TextExtractor._extract_from_pdf_with_metadata(file_path)
        elif file_extension == '.docx':
            return TextExtractor._extract_from_docx_with_metadata(file_path)
        elif file_extension == '.txt':
            return TextExtractor._extract_from_txt_with_metadata(file_path)
        else:
            raise ValueError(f"Неподдерживаемый формат файла: {file_extension}")
    
    @staticmethod
    def _extract_from_pdf(file_path: str) -> str:
        """Извлекает текст из PDF файла"""
        text = ""
        try:
            doc = fitz.open(file_path)
            for page in doc:
                text += page.get_text()
            doc.close()
        except Exception as e:
            raise Exception(f"Ошибка при извлечении текста из PDF: {str(e)}")
        return text
    
    @staticmethod
    def _extract_from_pdf_with_metadata(file_path: str) -> Dict[str, Any]:
        """Извлекает текст из PDF файла с метаданными и fallback на ИИ"""
        try:
            # Сначала пытаемся стандартным способом
            doc = fitz.open(file_path)
            text = ""
            pages_data = []
            total_tables = 0
            has_text_content = False
            
            for page_num, page in enumerate(doc):
                page_text = page.get_text()
                if page_text.strip():
                    has_text_content = True
                
                page_data = {
                    "page_number": page_num + 1,
                    "text": page_text,
                    "word_count": len(page_text.split()),
                    "char_count": len(page_text),
                    "tables": []
                }
                
                # Пытаемся извлечь таблицы (если метод доступен)
                try:
                    if hasattr(page, 'get_tables'):
                        tables = page.get_tables()
                        if tables:
                            total_tables += len(tables)
                            for table_idx, table in enumerate(tables):
                                table_text = ""
                                for row in table:
                                    table_text += " | ".join([str(cell) if cell else "" for cell in row]) + "\n"
                                page_data["tables"].append({
                                    "table_index": table_idx + 1,
                                    "text": table_text.strip(),
                                    "rows": len(table)
                                })
                except Exception:
                    # Если метод get_tables недоступен, пропускаем
                    pass
                
                pages_data.append(page_data)
                text += page_text + "\n"
            
            doc.close()
            
            # Если текста мало или нет, используем ИИ-конвертацию
            if not has_text_content or len(text.strip()) < 50:
                try:
                    from services.ai_pdf_converter import AIPDFConverter
                    ai_converter = AIPDFConverter()
                    ai_result = ai_converter.extract_text_with_ai_fallback(file_path)
                    
                    return {
                        "text": ai_result["text"],
                        "metadata": {
                            **ai_result["metadata"],
                            "conversion_method": "ai_fallback",
                            "original_attempt": {
                                "total_pages": len(pages_data),
                                "total_tables": total_tables,
                                "total_words": len(text.split()),
                                "total_chars": len(text),
                                "pages": pages_data
                            }
                        }
                    }
                except Exception as ai_error:
                    # Если ИИ-конвертация не удалась, возвращаем стандартный результат
                    return {
                        "text": text.strip(),
                        "metadata": {
                            "total_pages": len(pages_data),
                            "total_tables": total_tables,
                            "total_words": len(text.split()),
                            "total_chars": len(text),
                            "pages": pages_data,
                            "conversion_method": "standard",
                            "ai_fallback_failed": str(ai_error)
                        }
                    }
            
            return {
                "text": text.strip(),
                "metadata": {
                    "total_pages": len(pages_data),
                    "total_tables": total_tables,
                    "total_words": len(text.split()),
                    "total_chars": len(text),
                    "pages": pages_data,
                    "conversion_method": "standard"
                }
            }
            
        except Exception as e:
            raise Exception(f"Ошибка при извлечении текста из PDF: {str(e)}")
    
    @staticmethod
    def _extract_from_pdf_structured(file_path: str) -> str:
        """Извлекает текст из PDF с сохранением структуры страниц"""
        text = ""
        try:
            doc = fitz.open(file_path)
            for page_num, page in enumerate(doc):
                # Добавляем заголовок страницы
                text += f"\n{'='*60}\n"
                text += f"СТРАНИЦА {page_num + 1}\n"
                text += f"{'='*60}\n\n"
                
                # Извлекаем текст страницы
                page_text = page.get_text()
                
                # Разбиваем на абзацы и добавляем нумерацию
                paragraphs = [p.strip() for p in page_text.split('\n\n') if p.strip()]
                for i, paragraph in enumerate(paragraphs, 1):
                    if len(paragraph) > 10:  # Игнорируем очень короткие строки
                        text += f"Абзац {i}: {paragraph}\n\n"
                
                # Пытаемся извлечь таблицы (если метод доступен)
                try:
                    if hasattr(page, 'get_tables'):
                        tables = page.get_tables()
                        if tables:
                            text += f"--- ТАБЛИЦЫ НА СТРАНИЦЕ {page_num + 1} ---\n"
                            for table_idx, table in enumerate(tables, 1):
                                text += f"\nТаблица {table_idx}:\n"
                                text += "-" * 40 + "\n"
                                for row in table:
                                    text += " | ".join([str(cell) if cell else "" for cell in row]) + "\n"
                                text += "-" * 40 + "\n"
                except Exception:
                    # Если метод get_tables недоступен, пропускаем
                    pass
                
                text += "\n"
            doc.close()
        except Exception as e:
            raise Exception(f"Ошибка при извлечении текста из PDF: {str(e)}")
        return text
    
    @staticmethod
    def _extract_from_docx(file_path: str) -> str:
        """Извлекает текст из DOCX файла"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            raise Exception(f"Ошибка при извлечении текста из DOCX: {str(e)}")
    
    @staticmethod
    def _extract_from_docx_with_metadata(file_path: str) -> Dict[str, Any]:
        """Извлекает текст из DOCX файла с метаданными"""
        try:
            doc = Document(file_path)
            text = ""
            paragraphs_data = []
            
            for para_idx, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text
                if para_text.strip():
                    text += para_text + "\n"
                    paragraphs_data.append({
                        "index": para_idx + 1,
                        "text": para_text,
                        "word_count": len(para_text.split()),
                        "style": paragraph.style.name if paragraph.style else "Normal"
                    })
            
            return {
                "text": text.strip(),
                "metadata": {
                    "total_paragraphs": len(paragraphs_data),
                    "total_words": len(text.split()),
                    "paragraphs": paragraphs_data
                }
            }
        except Exception as e:
            raise Exception(f"Ошибка при извлечении текста из DOCX: {str(e)}")
    
    @staticmethod
    def _extract_from_txt(file_path: str) -> str:
        """Извлекает текст из TXT файла"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            raise Exception(f"Ошибка при извлечении текста из TXT: {str(e)}")
    
    @staticmethod
    def _extract_from_txt_with_metadata(file_path: str) -> Dict[str, Any]:
        """Извлекает текст из TXT файла с метаданными"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            lines = text.split('\n')
            paragraphs = [line.strip() for line in lines if line.strip()]
            
            return {
                "text": text,
                "metadata": {
                    "total_lines": len(lines),
                    "total_paragraphs": len(paragraphs),
                    "total_words": len(text.split()),
                    "total_chars": len(text)
                }
            }
        except Exception as e:
            raise Exception(f"Ошибка при извлечении текста из TXT: {str(e)}")
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """Разбивает текст на чанки с перекрытием"""
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Если это не последний чанк, ищем ближайший пробел для разрыва
            if end < len(text):
                # Ищем последний пробел в пределах чанка
                last_space = text.rfind(' ', start, end)
                if last_space > start:
                    end = last_space
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Следующий чанк начинается с учетом перекрытия
            start = end - overlap
            if start >= len(text):
                break
        
        return chunks
    
    @staticmethod
    def chunk_by_pages(text: str) -> List[str]:
        """Разбивает текст по страницам (для PDF с маркерами страниц)"""
        pages = text.split("СТРАНИЦА")
        chunks = []
        for page in pages:
            if page.strip():
                chunks.append(f"СТРАНИЦА{page.strip()}")
        return chunks
    
    @staticmethod
    def chunk_by_paragraphs(text: str, max_chunk_size: int = 1000) -> List[str]:
        """Разбивает текст по абзацам с ограничением размера"""
        paragraphs = text.split('\n\n')
        chunks = []
        current_chunk = ""
        
        for paragraph in paragraphs:
            if len(current_chunk) + len(paragraph) > max_chunk_size and current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = paragraph
            else:
                current_chunk += "\n\n" + paragraph if current_chunk else paragraph
        
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks 